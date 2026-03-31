from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

def heading(text, level=1):
    doc.add_heading(text, level=level)

def segment(time, title, screen, voiceover):
    p = doc.add_paragraph()
    r = p.add_run(f"{time}  |  {title}")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(43, 87, 151)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)

    ps = doc.add_paragraph()
    rs = ps.add_run("SCREEN: ")
    rs.bold = True
    rs.font.size = Pt(9)
    rs.font.color.rgb = RGBColor(100, 100, 100)
    rs2 = ps.add_run(screen)
    rs2.font.size = Pt(9)
    rs2.font.color.rgb = RGBColor(100, 100, 100)
    ps.paragraph_format.space_after = Pt(4)

    pv = doc.add_paragraph()
    rv = pv.add_run(voiceover)
    rv.font.size = Pt(11)
    pv.paragraph_format.space_after = Pt(10)
    pv.paragraph_format.left_indent = Cm(0.3)

heading("Demo Video Script — 5:00", level=1)

p = doc.add_paragraph()
r = p.add_run("Hedera Hello Future Apex Hackathon 2026 — Theme 3: Sustainability")
r.italic = True
r.font.size = Pt(10)

heading("Pre-Recording Setup", level=2)
setup = doc.add_paragraph()
setup.add_run("1. Open hedera-hello-future.vercel.app (clean browser, no bookmarks bar)\n"
              "2. Open HashScan testnet tab — topic 0.0.8217610\n"
              "3. Open HashScan testnet tab — token 0.0.8217620\n"
              "4. Recording: 1920x1080, no notifications\n"
              "5. Have GitHub repo open in another tab").font.size = Pt(10)

doc.add_paragraph()
heading("Script", level=2)

segment(
    "0:00–0:05", "Title Card",
    "RAEIS hero image from README",
    "(silence or brief music)"
)

segment(
    "0:05–0:35", "Problem Statement",
    "Navigate to /actions page. Slowly scroll through the 46 environmental actions. Click one action to show detail view — protocol name, certification, tCO2e, geography, source treasury account.",
    "Hedera has something most chains don't — real environmental data. Six Guardian-based platforms have tokenized carbon credits and renewable energy certificates on mainnet. DOVU, Tolam Earth, Capturiant, OrbexCO2, Global Carbon Registry, TYMLEZ. 46 verified environmental actions across 9 treasury accounts. But right now this data is locked inside individual HTS tokens. Each platform uses its own memo format, its own methodology, its own certification standard. There is no common language. If you are an AI agent trying to understand the environmental state of a bioregion, you have to reverse-engineer each platform separately. That is the problem we solve."
)

segment(
    "0:35–1:15", "RAEIS Three-Layer Standard",
    "Navigate to /publish page. Show the full transaction log — methodology topic at top, then bioregion feeds, then RAVA NFTs at bottom. Hover over entries to highlight transaction IDs. Click a HashScan link on the methodology topic.",
    "RAEIS — the Regen Atlas Environmental Intelligence Standard. We read all 46 actions from Hedera mainnet through Mirror Node. We trace each token's provenance by parsing treasury account memos — DOVU uses a colon-delimited format with topic IDs, Tolam embeds topics directly, Capturiant uses IPFS CIDs. We classify each action by certification tier and value it using the EPA Social Cost of Carbon at $51 to $190 per ton of CO2 equivalent. Then we publish a three-layer standard back to Hedera testnet. Layer 1 is a methodology topic on HCS — the machine-readable rulebook that defines how valuations work, what certification tiers mean, and what capabilities agents need to implement. Layer 2 is seven bioregional intelligence feeds, one HCS topic per bioregion, each carrying aggregated carbon data, service value ranges, and structured agent directives — VERIFY, BOUNTY, ALERT — telling machines what to do with the data. Layer 3 is the RAVA NFT collection on HTS — 46 verification attestation NFTs, one per environmental action."
)

segment(
    "1:15–1:55", "HashScan Deep Dive",
    "Switch to HashScan tab showing topic 0.0.8217610. Expand a message to show the methodology JSON. Then navigate to bioregion feed 0.0.8217612 — show the structured bioregion intelligence message with agent directives. Then switch to token tab showing RAVA collection 0.0.8217620 — scroll through NFT serials.",
    "Everything is onchain and verifiable. Here is the methodology topic on HashScan — you can see the full JSON defining carbon price ranges, the trust hierarchy, and the agent interface specification. This is the bioregion feed for Western European Broadleaf Forests — 7 actions, over 15,000 tonnes of CO2 equivalent, and three agent directives: verify the carbon tonnage, post a bounty for ground-truth data, and alert on an infinite gap factor where ecosystem service value dwarfs the market price. And here is the RAVA collection — 46 NFTs, each with metadata linking back to the specific environmental action it verifies. Three native Hedera services — Mirror Node for reads, HCS for ordered consensus messaging, HTS for tokenized attestations. No subgraph indexing. No external infrastructure. Just Hedera."
)

segment(
    "1:55–2:50", "OpenClaw Agent Network",
    "Navigate to HashScan topic 0.0.8218356 (Impact Scout). Show the OpportunityReport JSON — ranked bioregions with scores. Then navigate to topic 0.0.8218357 (Due Diligence). Show a DueDiligenceReport with PASS/CAUTION/FAIL verdicts. Zoom into a CAUTION verdict to show the reasoning.",
    "This is the OpenClaw bounty submission. Two autonomous agents coordinate entirely through Hedera Consensus Service — no shared database, no direct API calls between them. The Impact Opportunity Scout reads all seven bioregion feeds via Mirror Node, scores each bioregion on four axes — certification strength at 25%, carbon tonnage at 25%, data coverage at 15%, and market gap at 35% — and posts a ranked OpportunityReport to its own HCS topic. The Due Diligence Agent then discovers this report by reading that HCS topic through Mirror Node — exactly the way any third-party agent would find it. It takes each referenced token ID and cross-verifies it against Hedera mainnet. Does the token exist? Does the supply match? Is there a Guardian topic ID in the memo? It posts structured verdicts back to a second HCS topic. The results were real, not staged. 10 tokens passed full verification. 5 were flagged CAUTION — they exist on mainnet but their memos lack Guardian topic references, meaning provenance cannot be traced back to an MRV policy. That is a genuine data quality finding that came directly from the agent pipeline."
)

segment(
    "2:50–3:30", "Composability + Why It Matters",
    "Return to /publish page showing the full transaction log. Then switch to /explore map — zoom out to show global bioregion distribution. Pan across bioregion markers.",
    "What makes this different from a typical hackathon agent demo is composability. Everything the Scout and Diligence agents produce is published to public HCS topics. A third agent — say a capital allocation agent — could read both the Scout's opportunity rankings and the Diligence verdicts, then route funding to verified high-impact bioregions. A ground-truth agent could read the BOUNTY directives from bioregion feeds and dispatch human verifiers to collect soil samples or photograph restoration sites. None of these agents need permission from us. They just read public topics via Mirror Node. HCS gives you ordered, timestamped consensus messaging — something no EVM chain offers natively. For AI agent coordination, this is a genuine infrastructure advantage. Trust comes from consensus, not from access control."
)

segment(
    "3:30–4:10", "Valuation Engine + Intelligence Dashboard",
    "Navigate to /intelligence page. Show the valuation dashboard — SCC ranges, certification tier breakdown, ecological impact gaps. Click into a specific bioregion to show per-action valuations. Show the trust hierarchy visualization.",
    "The valuation engine uses the EPA Social Cost of Carbon from their 2024 technical update — $51 per ton at the low end, $190 at the high end, with a central estimate around $120. Each action is weighted by its certification tier. Verra VCS and Gold Standard carry full weight. Platform-specific dMRV standards like DOVU carry 0.7. Bare HTS tokens with no Guardian provenance carry 0.3. The ecological impact gap — the ratio between what these ecosystem services are worth according to peer-reviewed literature and what the carbon market actually prices them at — is typically 10x or higher. That gap is the investment signal. RAEIS makes it machine-readable."
)

segment(
    "4:10–4:40", "Guardian Integration + Research",
    "Show the README section on Guardian integration — the platform table with treasury accounts, action counts, and certifications. Briefly show a code snippet from the connector if screen time allows.",
    "The data pipeline was the hardest part. We identified 155 Hedera tokens across Guardian platforms, filtered to 115 after removing dead projects and illegitimate tokens — EcoGuard, RECDeFi, and BCarbon turned out to be inactive or empty. 46 actions survived geography and legitimacy filtering. Each platform required custom memo parsing. DOVU, Tolam, Capturiant, OrbexCO2, GCR, and TYMLEZ all use different formats. A standardized Guardian token memo format would make this kind of cross-platform intelligence far easier to build."
)

segment(
    "4:40–4:55", "Tech Stack + Track",
    "Show GitHub repo briefly — file structure, README, license. Return to live demo.",
    "Theme 3: Sustainability. React, TypeScript, Hashgraph SDK, Mirror Node REST API, OpenClaw agent framework. 67 testnet transactions. 46 mainnet Guardian actions ingested. 6 platforms. 7 bioregions. Two autonomous agents. MIT license. Live at hedera-hello-future.vercel.app."
)

segment(
    "4:55–5:00", "Close",
    "Live demo map view with bioregion markers. GitHub URL overlay fades in.",
    "Regen Atlas. Environmental intelligence for AI agents. Built on Hedera."
)

doc.add_page_break()
heading("Post-Recording Checklist", level=2)

checks = [
    "/actions page shown with real Guardian data (click into detail)",
    "/publish page with transaction log + HashScan link clicked",
    "HashScan deep dive: methodology topic JSON, bioregion feed with directives, RAVA NFT collection",
    "Scout HCS topic shown with OpportunityReport (0.0.8218356)",
    "Diligence HCS topic shown with verdicts including CAUTION finding (0.0.8218357)",
    "Composability argument made (third-party agents, no permission needed)",
    "/intelligence valuation dashboard with SCC ranges + impact gaps",
    "Guardian platform table + research effort acknowledged",
    "Three Hedera services named (Mirror Node, HCS, HTS)",
    "OpenClaw bounty connection explicit",
    "Runtime <= 5:00",
    "Upload to YouTube (unlisted), get link",
]

for c in checks:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(c)
    r.font.size = Pt(10)

output = "/Users/pat/Desktop/1_projects/hedera-hello-future/hedera-demo-script.docx"
doc.save(output)
print(f"Saved to {output}")
